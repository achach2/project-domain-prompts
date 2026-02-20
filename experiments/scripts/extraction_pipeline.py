import os
import re
import docx

def rename_files_in_directory(directory_path):
    """
    Renames all files in the specified directory by replacing spaces in filenames with underscores.

    Parameters:
    directory_path (str): The path to the directory containing the files to be renamed.

    Notes:
    - Only regular files are renamed; subdirectories are ignored.
    - If a filename does not contain any spaces, it will not be renamed.
    - This operation is not recursive; files in subdirectories are not affected.
    """
    try:
        for filename in os.listdir(directory_path):
            old_path = os.path.join(directory_path, filename)
            if os.path.isfile(old_path):
                new_filename = filename.replace(' ', '_')
                new_path = os.path.join(directory_path, new_filename)
                if old_path != new_path:
                    os.rename(old_path, new_path)
                    print(f'Renamed: "{filename}" -> "{new_filename}"')
    except Exception as e:
        print(f"Error: {e}")
        
        
def extract_headings_and_text(docx_file):
    """
    Extracts headings and their associated paragraph text from a .docx file.

    Parameters:
    docx_file (str): Path to the input .docx file.

    Returns:
    list of [str, str]: A list of pairs where each pair contains a heading and its combined paragraph text.
    """
    doc = docx.Document(docx_file)
    result, current_heading, current_text = [], None, []

    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            if current_heading:
                result.append([current_heading, ' '.join(current_text)])
            current_heading = para.text
            current_text = []
        else:
            if para.text.strip():
                current_text.append(para.text.strip())

    if current_heading:
        result.append([current_heading, ' '.join(current_text)])

    return result
    
def extract_sections_and_section_numbers(docx_file):
    """
    Extracts section numbers (e.g., '1.1', '2.3') and section titles from a .docx file.

    Parameters:
    docx_file (str): Path to the input .docx file.

    Returns:
    dict: A dictionary mapping section titles to section numbers.
    """
    doc = docx.Document(docx_file)
    result, current_section, current_heading, current_text = [], None, None, []

    section_regex = re.compile(r'^(\d+(\.\d+)*)\s+(.+)')

    for para in doc.paragraphs:
        match = section_regex.match(para.text)
        if match:
            if current_section:
                result.append([current_section, current_heading, ' '.join(current_text)])
            current_section, current_heading, current_text = match.group(1), match.group(3), []
        else:
            if para.text.strip():
                current_text.append(para.text.strip())

    if current_section:
        result.append([current_section, current_heading, ' '.join(current_text)])

    sections = {}
    for section in result:
        if not section[1]:  # Skip if heading is empty
            section_data = re.sub(r'\t+| +', ' ', section[0])
            section_number = section_data.split(" ")[0]
            section_name = section_data.lstrip(f"{section_number} ")
            sections[section_name] = section_number

    return sections

def create_section_knowledge_database(file_path):
    """
    Creates a structured knowledge database from a single .docx file using headings and section metadata.

    Parameters:
    file_path (str): Path to the .docx file.

    Returns:
    tuple: 
        - list of dict: Structured data entries with metadata and text.
        - str: The filename of the processed document.
    """
    file_name = file_path.split("/")[-1].rstrip(".docx")
    section_metadata = extract_sections_and_section_numbers(file_path)
    section_number_to_header = {v: k for k, v in section_metadata.items()}
    text_data = extract_headings_and_text(file_path)
    
    database, sections_list, section_numbers_list = [], list(section_metadata.keys()), list(section_number_to_header.keys())

    for text in text_data:
        if text[1] and text[0] in sections_list:
            index = section_metadata[text[0]].rfind(".")
            parent_section_number = section_metadata[text[0]][:index]
            parent_section_name = section_number_to_header[parent_section_number] if parent_section_number in section_numbers_list else ""
            database.append({
                'filename': file_name,
                'Text': text[1],
                'Subsection Name': text[0],
                'Subsection Number': section_metadata[text[0]],
                'Section Name': parent_section_name,
                'Section Number': parent_section_number
            })

    return database, file_name
    
def create_section_knowledge_database_byfolder(folder_path):
    """
    Creates a section-level knowledge database from all .docx files in a folder.

    Parameters:
    folder_path (str): Path to the folder containing .docx files.

    Returns:
    list of dict: A combined list of section knowledge entries from all processed files.
    """
    folder_database = []
    files = [f for f in os.listdir(folder_path) if f.endswith(".docx")]

    for file in files:
        file_path = f'{folder_path}/{file}'
        database, file_name = create_section_knowledge_database(file_path)
        folder_database.extend(database)
        print(f"Completed file: {file_path}")

    return folder_database